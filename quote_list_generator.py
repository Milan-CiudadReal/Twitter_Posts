import docx
import random
import tweepy

doc = docx.Document("Inspirational Quotes.docx")
# print(len(doc.paragraphs))
# print(doc.paragraphs[152].text)

string = []
quotes_list = []
i = 128
quote = ""
while i < 3646:
    string.append(doc.paragraphs[i].text)
    i += 1
# print(string)
# print(string.index('–Charles Spurgeon'))

while '' in string:
    string.remove('')

for text in string:
    try:
        text == int(text)
        quotes_list.append(quote)
        # print(quote)
        quote = ""


    except ValueError:
        quote += f"{text} "



random.shuffle(quotes_list)
print(quotes_list[0])

# API_key = "qtYX8mxTdx4QLXX0zDYhMb73l"
#
# secret_key = "o8PGtmDS5r2tmTxbUpp6D2MfE2fkSQCqIOjGwaXo063wF5vwn8"
#
# access_token = "1259716202630430720-wjCQGdVGZziVjtgzBT9LaRh9KgtL55"
#
# secret_access_token = "FGmH7Bb5FNYk6YCMcMi3DUwZudYfzlN8uz23uBo3SU82V"
#
#
# auth = tweepy.OAuthHandler(API_key, secret_key)
# auth.set_access_token(access_token, secret_access_token)
#
# api = tweepy.API(auth)
# api.update_status("Russo's a pussy")

